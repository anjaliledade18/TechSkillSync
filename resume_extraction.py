import os
import re
import pdfplumber
from docx import Document
import spacy
from ner_model import ner_model

#Load spaCy model once
nlp = spacy.load("en_core_web_sm")

#=== Main function ===
def extract_skills_from_file_or_text(input_data, known_skills):
    if isinstance(input_data, str) and os.path.exists(input_data):
        lines = _parse_document(input_data)
        raw_text = "\n".join(lines)
    elif isinstance(input_data, str):
        raw_text = input_data
    elif isinstance(input_data, list):
        raw_text = "\n".join(input_data)
    else:
        raise ValueError("Input must be a file path, raw text string, or list of lines.")

    cleaned_text = _clean_text(raw_text)
    skills_found = _extract_skills(cleaned_text, set(known_skills))
    contextual_skills = _extract_contextual_skills(cleaned_text, set(known_skills))
    ner_skills = _extract_skills_from_ner(cleaned_text)
    if ner_skills:
        return list(set(skills_found + contextual_skills + ner_skills))
    return list(set(skills_found + contextual_skills))


#=== File Parsing ===
def _parse_document(file_path):
    ext = os.path.splitext(file_path)[-1].lower()
    if ext == ".docx":
        return _parse_word(file_path)
    elif ext == ".pdf":
        return _parse_pdf(file_path)
    else:
        raise ValueError("Unsupported file format: use .pdf or .docx")


def _parse_word(file_path):
    doc = Document(file_path)
    lines = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    lines.append(cell_text)
    return lines


def _parse_pdf(file_path):
    lines = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                lines.extend(line.strip() for line in page_text.split('\n') if line.strip())
    return lines


#=== Skill Extraction Logic ===
def _extract_skills(text, skill_set):
    doc = nlp(text)
    extracted = set()

    for token in doc:
        if token.text.lower() in skill_set:
            extracted.add(token.text.lower())

    for chunk in doc.noun_chunks:
        chunk_text = chunk.text.lower().strip()
        if chunk_text in skill_set:
            extracted.add(chunk_text)

    return list(extracted)


def _extract_contextual_skills(text, skill_set):
    return [skill.lower() for skill in skill_set if re.search(r'\b' + re.escape(skill) + r'\b', text, re.IGNORECASE)]


def _clean_text(text):
    return re.sub(r'\s+', ' ', text).lower().strip()

def _extract_skills_from_ner(text):
    ner_extracted_skills = []
    try:
        ner_extracted_skills = ner_model(text)
    except Exception as e:
        print(e)
    return ner_extracted_skills